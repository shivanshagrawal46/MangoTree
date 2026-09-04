"""Native text extraction — the free path that runs before vision OCR.

Most PDFs in this corpus are digitally generated (title letters, loan documents,
investor packages) and carry a real text layer. Extracting that layer is exact,
instant and free; rendering those same pages to images and asking a vision model
to read them would cost money to produce a *worse* transcription.

So the rule is: **take the text layer when it is genuinely there, and escalate to
vision only for the pages that need it** — scans, photos of documents, and pages
whose text layer is empty or suspiciously thin.

``page_needs_vision`` is deliberately conservative. A page wrongly sent to vision
costs a few cents; a page wrongly trusted enters the corpus with half its content
missing, and nothing downstream will ever notice.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from mangotree.core.logging import logger

#: A page with fewer characters than this is treated as empty/scanned.
MIN_CHARS_PER_PAGE = 120

#: Ratio of alphanumeric characters below which the "text" is probably
#: extraction garbage (broken embedded fonts produce exactly this).
MIN_ALNUM_RATIO = 0.55


@dataclass
class PageText:
    page: int
    text: str
    char_count: int
    needs_vision: bool
    reason: str = ""


@dataclass
class NativeExtract:
    path: str
    kind: str
    pages: List[PageText] = field(default_factory=list)
    text: str = ""
    errors: List[str] = field(default_factory=list)

    @property
    def pages_needing_vision(self) -> List[int]:
        return [p.page for p in self.pages if p.needs_vision]

    @property
    def has_usable_text(self) -> bool:
        return bool(self.text.strip())

    def as_dict(self) -> dict:
        return {
            "kind": self.kind,
            "page_count": len(self.pages),
            "char_count": len(self.text),
            "pages_needing_vision": self.pages_needing_vision,
            "errors": self.errors,
        }


def _alnum_ratio(text: str) -> float:
    stripped = re.sub(r"\s", "", text)
    if not stripped:
        return 0.0
    return sum(c.isalnum() or c in ".,;:$%()/-" for c in stripped) / len(stripped)


def page_needs_vision(text: str) -> Tuple[bool, str]:
    """Decide whether a page's native text can be trusted."""
    cleaned = (text or "").strip()
    if len(cleaned) < MIN_CHARS_PER_PAGE:
        return True, f"only {len(cleaned)} chars of native text (scanned or image-only)"
    ratio = _alnum_ratio(cleaned)
    if ratio < MIN_ALNUM_RATIO:
        return True, f"text layer looks corrupt (alnum ratio {ratio:.2f})"
    return False, ""


# --------------------------------------------------------------------------
def extract_pdf_text(path: Path) -> NativeExtract:
    result = NativeExtract(path=str(path), kind="pdf")
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(path)
        try:
            for index, page in enumerate(doc, start=1):
                text = page.get_text("text") or ""
                needs, reason = page_needs_vision(text)
                result.pages.append(
                    PageText(index, text, len(text.strip()), needs, reason)
                )
        finally:
            doc.close()
    except ImportError:
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            for index, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                needs, reason = page_needs_vision(text)
                result.pages.append(
                    PageText(index, text, len(text.strip()), needs, reason)
                )
        except Exception as exc:
            result.errors.append(f"{type(exc).__name__}: {exc}")
    except Exception as exc:
        result.errors.append(f"{type(exc).__name__}: {exc}")

    result.text = "\n\n".join(
        f"[page {p.page}]\n{p.text.strip()}" for p in result.pages if p.text.strip()
    )
    return result


def extract_docx(path: Path) -> NativeExtract:
    result = NativeExtract(path=str(path), kind="document")
    try:
        import docx

        document = docx.Document(str(path))
        parts: List[str] = [p.text for p in document.paragraphs if p.text.strip()]

        # Tables carry the numbers in these documents; losing them would drop
        # exactly the content the ledger cares about.
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        result.text = "\n".join(parts)
        result.pages = [PageText(1, result.text, len(result.text), False)]
    except Exception as exc:
        result.errors.append(f"{type(exc).__name__}: {exc}")
        logger.error("docx extraction failed for %s: %s", path.name, exc)
    return result


def extract_text_file(path: Path) -> NativeExtract:
    result = NativeExtract(path=str(path), kind="text")
    try:
        result.text = Path(path).read_text(encoding="utf-8", errors="replace")
        result.pages = [PageText(1, result.text, len(result.text), False)]
    except Exception as exc:
        result.errors.append(f"{type(exc).__name__}: {exc}")
    return result


def extract_native(path: Path) -> NativeExtract:
    """Dispatch native extraction by file type."""
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(path)
    if suffix in {".docx", ".doc"}:
        return extract_docx(path)
    if suffix in {".txt", ".md"}:
        return extract_text_file(path)
    result = NativeExtract(path=str(path), kind="unsupported")
    result.errors.append(f"no native extractor for {suffix}")
    return result
